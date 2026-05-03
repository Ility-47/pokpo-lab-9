from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"c:\university\8sem\покпо\lab7")
OUTPUT = ROOT / "lab9_report_ci_cd_code_quality.docx"


def read_text(path: Path) -> str:
    data = path.read_bytes()
    for enc in ("utf-8", "utf-16", "utf-16-le", "utf-16-be", "cp1251"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace").strip()


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_paragraph(doc: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_code_block(doc: Document, title: str, content: str) -> None:
    add_paragraph(doc, title, bold=True)
    for line in content.splitlines():
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.first_line_indent = Cm(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(11)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.font.size = Pt(14)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

title_lines = [
    "МИНОБРНАУКИ РОССИИ",
    "Федеральное государственное бюджетное образовательное учреждение высшего образования",
    "«Тульский государственный университет»",
    "Институт прикладной математики и компьютерных наук",
    "Кафедра вычислительной техники",
    "",
    "Отчет по лабораторной работе №9",
    "на тему «Тестирование в CI/CD и анализ качества кода»",
    "по дисциплине «Процессы обеспечения качества программного обеспечения»",
    "",
    "Выполнил: студент ______________________________",
    "Проверил: _____________________________________",
    "",
    "Тула 2026",
]
for line in title_lines:
    add_paragraph(doc, line, center=True)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

add_heading(doc, "Цель работы")
add_paragraph(
    doc,
    "Изучить принципы интеграции процесса тестирования в конвейер непрерывной интеграции и доставки (CI/CD), "
    "освоить инструменты статического и динамического анализа кода, а также научиться настраивать автоматическую "
    "генерацию отчетов о качестве и покрытии кода тестами.",
)

add_heading(doc, "Описание статического анализа (Задание 1)")
add_paragraph(
    doc,
    "В проект были интегрированы следующие инструменты статического анализа:",
)
add_paragraph(doc, "1) StyleCop.Analyzers — анализатор стиля кода для C# (NuGet-пакет).", bold=False)
add_paragraph(
    doc,
    "2) Встроенный анализатор .NET (Roslyn/CA правила) — проверка качества и потенциальных ошибок при сборке.",
)
add_paragraph(
    doc,
    "3) Lizard — анализ цикломатической сложности (CCN) и размеров функций. Используется для оценки сложности и "
    "выявления потенциально трудно тестируемых участков.",
)
add_paragraph(
    doc,
    "Обоснование выбора: инструменты не требуют отдельного сервера, легко подключаются в CI, дают повторяемые результаты "
    "и подходят для проектов на .NET и WinForms.",
)

complexity_path = ROOT / "artifacts_complexity.txt"
if complexity_path.exists():
    add_code_block(doc, "Пример отчета Lizard (цикломатическая сложность)", read_text(complexity_path))
else:
    add_paragraph(
        doc,
        "Отчет Lizard формируется в CI как артефакт (static-analysis-artifacts/complexity-report.txt).",
    )

add_heading(doc, "Описание CI/CD пайплайна (Задание 2)")
add_paragraph(
    doc,
    "CI/CD реализован на GitHub Actions. Пайплайн запускается при push в любую ветку и при pull request в main. "
    "Состоит из следующих jobs: static-analysis, unit-tests, integration-tests, coverage-report. "
    "Результаты тестов и отчеты публикуются как artifacts.",
)
add_code_block(doc, "Листинг .github/workflows/ci.yml", read_text(ROOT / ".github" / "workflows" / "ci.yml"))

add_paragraph(
    doc,
    "Дополнительно настроен security-анализ CodeQL, который выполняется при push/PR в main и по расписанию.",
)
add_code_block(doc, "Листинг .github/workflows/codeql.yml", read_text(ROOT / ".github" / "workflows" / "codeql.yml"))

add_paragraph(
    doc,
    "Branch protection на GitHub настраивается в Settings → Branches: включить требование успешного прохождения "
    "проверок перед merge (Require status checks to pass before merging) и выбрать проверки workflow.",
)

add_heading(doc, "Описание анализа покрытия и генерации отчетов (Задание 3)")
add_paragraph(
    doc,
    "Для измерения покрытия используется Coverlet Collector. Порог покрытия настроен в файле coverlet.runsettings "
    "(Threshold = 70%, ThresholdType = line, ThresholdStat = total).",
)
add_code_block(doc, "Листинг coverlet.runsettings", read_text(ROOT / "coverlet.runsettings"))

summary_path = ROOT / "coveragereport" / "Summary.txt"
if summary_path.exists():
    add_code_block(doc, "Сводка покрытия (ReportGenerator Summary.txt)", read_text(summary_path))
    add_paragraph(
        doc,
        "Анализ белых пятен: непокрыты UI-формы WinForms и CLI-часть приложения, так как текущие unit/integration "
        "сценарии фокусируются на бизнес-логике (MediaService/PasswordPolicy) и BDD шагах. Для увеличения покрытия "
        "можно добавить отдельные UI-тесты/интеграцию для WinForms форм или вынести логику из форм в сервисы.",
    )
else:
    add_paragraph(
        doc,
        "HTML-отчет покрытия генерируется job'ом coverage-report и прикладывается как artifact (coverage-report).",
    )

add_heading(doc, "Выводы")
add_paragraph(
    doc,
    "Интеграция статического анализа, запуска тестов и генерации отчетов покрытия в CI/CD позволяет обнаруживать "
    "ошибки и регрессии на ранних этапах, поддерживать единый стиль кода и получать наглядные метрики качества. "
    "Это повышает надежность продукта и снижает затраты на ручные проверки за счет автоматизации.",
)

add_heading(doc, "Ответы на контрольные вопросы")
answers = [
    "1) CI/CD — практика автоматизации сборки, тестирования и доставки/развертывания. Компоненты: pipeline, stages/jobs, runner, artifacts, cache.",
    "2) Continuous Delivery — код готов к релизу после CI, но деплой может быть вручную. Continuous Deployment — деплой выполняется автоматически после успешных проверок.",
    "3) Pipeline — общий конвейер; stage/job — задача или группа задач, выполняющая конкретный шаг (build/test/report).",
    "4) Автотесты в CI/CD ускоряют обнаружение дефектов и предотвращают попадание сломанного кода в main.",
    "5) Artifacts — результаты (отчеты/логи), которые сохраняются. Cache — кеш зависимостей для ускорения сборок.",
    "6) Статический анализ — анализ кода без выполнения, помогает найти нарушения стиля, потенциальные ошибки, smells, уязвимости.",
    "7) Статический анализ не исполняет программу, динамический (тестирование) проверяет поведение на запуске.",
    "8) Code smell — признак потенциальной проблемы (длинные методы, дублирование, магические числа).",
    "9) Цикломатическая сложность — количество независимых путей в коде; высокая сложность усложняет тестирование и поддержку.",
    "10) Инструменты: StyleCop.Analyzers, SonarQube/SonarCloud, Roslyn Analyzers, ReSharper, CodeQL.",
    "11) Code coverage — доля кода, выполняемая тестами; виды: line coverage, branch coverage и др.",
    "12) Инструменты покрытия: Coverlet/dotnet-coverage (.NET), JaCoCo (Java), coverage.py/pytest-cov (Python).",
    "13) 100% coverage не гарантирует отсутствие ошибок: тесты могут не проверять корректность логики и граничные случаи.",
    "14) Allure Reports — генерация наглядных отчетов о тестах (шаги, вложения, статистика).",
    "15) Публикация отчетов в Pages: генерировать статический HTML и публиковать в ветку gh-pages через GitHub Actions.",
    "16) Branch protection: Settings → Branches → Add rule → Require status checks to pass before merging.",
    "17) CodeQL — статический анализ безопасности от GitHub, ищет уязвимости и небезопасные паттерны.",
    "18) Секреты нужно хранить в GitHub Secrets/Variables и не печатать в логах.",
    "19) allow_failure (GitLab) — разрешает задаче падать, не ломая весь пайплайн (для необязательных проверок).",
    "20) Ускорение CI/CD: кеширование зависимостей, параллелизм (matrix), разделение тестов, уменьшение лишних шагов.",
]
for a in answers:
    add_paragraph(doc, a)

doc.save(OUTPUT)
print(f"Created: {OUTPUT}")
